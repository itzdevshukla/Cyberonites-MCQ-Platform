"""
DOCX Question Parser.
Robustly parses DOCX files into structured question data.
Supports paragraphs, Word tables, numbered questions, and flexible option/answer formats.
"""
import re
import logging
from docx import Document

logger = logging.getLogger('quiz')


class DocxQuestionParser:
    """
    Super-flexible parser for MCQ questions in Word (.docx) files.
    """

    # Field patterns (case-insensitive & multi-format)
    PATTERNS = {
        'question': re.compile(
            r'^(?:question|q\d*|\d+)\s*[:\.\)]\s*(.+)', re.IGNORECASE
        ),
        'option_a': re.compile(
            r'^(?:option\s*a|\(a\)|a)\s*[:\.\)]\s*(.+)', re.IGNORECASE
        ),
        'option_b': re.compile(
            r'^(?:option\s*b|\(b\)|b)\s*[:\.\)]\s*(.+)', re.IGNORECASE
        ),
        'option_c': re.compile(
            r'^(?:option\s*c|\(c\)|c)\s*[:\.\)]\s*(.+)', re.IGNORECASE
        ),
        'option_d': re.compile(
            r'^(?:option\s*d|\(d\)|d)\s*[:\.\)]\s*(.+)', re.IGNORECASE
        ),
        'correct': re.compile(
            r'^(?:correct\s*(?:option|answer|ans)?|answer|ans|right\s*option)\s*[:\.]\s*(.+)', re.IGNORECASE
        ),
        'description': re.compile(
            r'^(?:description|desc|explanation|exp|note)\s*[:\.]\s*(.+)', re.IGNORECASE
        ),
        'topic': re.compile(
            r'^(?:topic|subject|category)\s*[:\.]\s*(.+)', re.IGNORECASE
        ),
        'difficulty': re.compile(
            r'^(?:difficulty|level)\s*[:\.]\s*(.+)', re.IGNORECASE
        ),
    }

    VALID_DIFFICULTIES = {'EASY', 'MEDIUM', 'HARD'}

    def __init__(self, file_path):
        self.file_path = file_path
        self.questions = []
        self.errors = []

    def parse(self):
        """Parse DOCX file paragraphs and tables."""
        try:
            doc = Document(self.file_path)
        except Exception as e:
            self.errors.append(f"Failed to open document: {str(e)}")
            return self.questions

        # Collect text lines from both paragraphs and tables (split by newlines)
        lines = []
        for para in doc.paragraphs:
            if para.text:
                for line in para.text.splitlines():
                    txt = line.strip()
                    if txt:
                        lines.append(txt)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        for line in cell.text.splitlines():
                            txt = line.strip()
                            if txt and txt not in lines:
                                lines.append(txt)


        current = self._empty_question()
        question_num = 0

        for text in lines:
            matched = False
            for field, pattern in self.PATTERNS.items():
                match = pattern.match(text)
                if match:
                    value = match.group(1).strip()

                    # Save current question when encountering a new question line
                    if field == 'question' and current['question']:
                        question_num += 1
                        self._validate_and_add(current, question_num)
                        current = self._empty_question()

                    current[field] = value
                    matched = True
                    break

            if not matched and current['question']:
                # Continuation text if no pattern matched
                current['question'] += ' ' + text

        # Add the final question
        if current['question']:
            question_num += 1
            self._validate_and_add(current, question_num)

        logger.info(f"Parsed {len(self.questions)} questions from DOCX, {len(self.errors)} errors")
        return self.questions

    def _empty_question(self):
        return {
            'question': '',
            'option_a': '',
            'option_b': '',
            'option_c': '',
            'option_d': '',
            'correct': '',
            'description': '',
            'topic': 'General',
            'difficulty': 'MEDIUM',
        }

    def _validate_and_add(self, q, num):
        """Validate a parsed question block and record errors if any."""
        errors = []

        if not q['question']:
            errors.append(f"Q{num}: Missing question text")
        if not q['option_a']:
            errors.append(f"Q{num}: Missing Option A")
        if not q['option_b']:
            errors.append(f"Q{num}: Missing Option B")
        if not q['option_c']:
            errors.append(f"Q{num}: Missing Option C")
        if not q['option_d']:
            errors.append(f"Q{num}: Missing Option D")

        # Extract correct option letter (A, B, C, D)
        correct_raw = q['correct'].upper().strip()
        correct = ''

        # Search for first occurrence of A, B, C, or D
        match_letter = re.search(r'\b([A-D])\b', correct_raw)
        if match_letter:
            correct = match_letter.group(1)
        elif correct_raw and correct_raw[0] in ['A', 'B', 'C', 'D']:
            correct = correct_raw[0]

        if not correct:
            errors.append(f"Q{num}: Invalid correct option '{q['correct']}'. Must specify A, B, C, or D.")

        difficulty = q['difficulty'].upper().strip()
        if difficulty not in self.VALID_DIFFICULTIES:
            difficulty = 'MEDIUM'

        if errors:
            self.errors.extend(errors)
        else:
            self.questions.append({
                'question': q['question'],
                'options': {
                    'A': q['option_a'],
                    'B': q['option_b'],
                    'C': q['option_c'],
                    'D': q['option_d'],
                },
                'correct': correct,
                'description': q['description'],
                'topic': q['topic'] or 'General',
                'difficulty': difficulty,
            })
