"""
Views for question upload and template download.
"""
import logging
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse

from accounts.decorators import admin_required
from quiz.models import Quiz, Question, Option
from .models import QuestionUpload
from .parsers import DocxQuestionParser

logger = logging.getLogger('quiz')


@admin_required
def upload_questions(request, quiz_id):
    """Handle DOCX file upload and parsing."""
    quiz = get_object_or_404(Quiz, id=quiz_id)

    if request.method == 'POST':
        uploaded_file = request.FILES.get('docx_file')
        if not uploaded_file:
            messages.error(request, "Please select a DOCX file.")
            return redirect('dashboard:upload_questions', quiz_id=quiz_id)

        if not uploaded_file.name.endswith('.docx'):
            messages.error(request, "Only .docx files are supported.")
            return redirect('dashboard:upload_questions', quiz_id=quiz_id)

        # Create upload record
        upload = QuestionUpload.objects.create(
            file=uploaded_file,
            quiz=quiz,
            uploaded_by=request.user,
        )

        # Parse DOCX
        parser = DocxQuestionParser(upload.file.path)
        parsed_questions = parser.parse()

        upload.total_parsed = len(parsed_questions) + len(parser.errors)

        # Save questions to database
        saved_count = 0
        existing_count = quiz.questions.count()

        for i, q_data in enumerate(parsed_questions):
            try:
                question = Question.objects.create(
                    quiz=quiz,
                    text=q_data['question'],
                    topic=q_data['topic'],
                    difficulty=q_data['difficulty'],
                    description=q_data['description'],
                    order=existing_count + i + 1,
                    marks=1,
                    negative_marks=0,
                )

                for label, text in q_data['options'].items():
                    Option.objects.create(
                        question=question,
                        label=label,
                        text=text,
                        is_correct=(label == q_data['correct']),
                    )
                saved_count += 1
            except Exception as e:
                parser.errors.append(f"Failed to save question {i + 1}: {str(e)}")
                logger.error(f"Failed to save question: {e}")

        # Update upload record
        upload.total_saved = saved_count
        upload.error_log = '\n'.join(parser.errors) if parser.errors else ''
        upload.status = 'SUCCESS' if not parser.errors else ('PARTIAL' if saved_count > 0 else 'FAILED')
        upload.save()

        # Recalculate quiz total marks
        quiz.recalculate_total_marks()

        if saved_count > 0:
            messages.success(request, f"Successfully imported {saved_count} questions!")
        if parser.errors:
            messages.warning(request, f"{len(parser.errors)} errors occurred during import.")

        return redirect('dashboard:question_list', quiz_id=quiz_id)

    context = {
        'quiz': quiz,
        'uploads': QuestionUpload.objects.filter(quiz=quiz)[:10],
    }
    return render(request, 'dashboard/admin/upload_questions.html', context)


@admin_required
def download_template(request):
    """Serve a sample DOCX template for question upload."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading('Question Upload Template', 0)
    doc.add_paragraph('Fill in questions below following this exact format. Do not change the labels.')
    doc.add_paragraph('')

    # Sample question
    doc.add_paragraph('Question: What is the capital of France?')
    doc.add_paragraph('Option A: Berlin')
    doc.add_paragraph('Option B: Paris')
    doc.add_paragraph('Option C: Madrid')
    doc.add_paragraph('Option D: Rome')
    doc.add_paragraph('Correct Option: B')
    doc.add_paragraph('Description: Paris is the capital and most populous city of France.')
    doc.add_paragraph('Topic: Geography')
    doc.add_paragraph('Difficulty: Easy')
    doc.add_paragraph('')

    # Second sample
    doc.add_paragraph('Question: Which data structure uses LIFO?')
    doc.add_paragraph('Option A: Queue')
    doc.add_paragraph('Option B: Stack')
    doc.add_paragraph('Option C: Array')
    doc.add_paragraph('Option D: Linked List')
    doc.add_paragraph('Correct Option: B')
    doc.add_paragraph('Description: Stack follows Last In First Out (LIFO) principle.')
    doc.add_paragraph('Topic: Data Structures')
    doc.add_paragraph('Difficulty: Medium')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=question_template.docx'
    doc.save(response)
    return response
